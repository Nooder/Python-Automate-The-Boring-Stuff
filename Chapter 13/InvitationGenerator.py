#! python3
# Chapter 13 Project - Generate invitations using names from guests.txt

import docx, os, sys, logging

# Setup
txtFile = open("guests.txt")
names = txtFile.readlines()
doc = docx.Document()
# Make the docx for each name
for name in names:
    doc.add_heading("You're invited!", 0)
    doc.add_paragraph(name)
    doc.add_paragraph("Regards: Mr. Monkey =)")
    doc.paragraphs[names.index(name) * 3 - 1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

doc.save("All_invitations.docx")
print("Done.")